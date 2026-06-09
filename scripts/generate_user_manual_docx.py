#!/usr/bin/env python3
"""Generate the Japanese Windows user manual DOCX from its Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "WINDOWS_USER_MANUAL_JA.md"
OUTPUT = ROOT / "docs" / "AppWatchdog_Windows_User_Manual_JA.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
TABLE_WIDTHS_DXA = {
    2: [3000, 6360],
    3: [1750, 2850, 4760],
}


def set_run_font(run, name: str = "Yu Gothic", size: float | None = None,
                 color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 28, DARK_BLUE, 0, 8),
        ("Subtitle", 13, MUTED, 0, 18),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, DARK_BLUE, 10, 5),
    ):
        style = doc.styles[name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(10)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_page_number(paragraph) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MUTED)
    field_run = paragraph.add_run()
    set_run_font(field_run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_char1)
    field_run._r.append(instr)
    field_run._r.append(fld_char2)


def configure_section(section) -> None:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "AppWatchdog | Windows 操作マニュアル"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in hp.runs:
        set_run_font(run, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("AppWatchdog")
    set_run_font(run, size=30, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Windows 操作マニュアル")
    set_run_font(run, size=18, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("ダウンロード・初期設定・自動起動・日常操作・トラブル対応")
    set_run_font(run, size=11, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "推奨: Windowsログオン後も監視を自動開始する場合は、"
        "スタートアップフォルダーへAppWatchdogのショートカットを登録します。"
    )
    set_run_font(run, size=11, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("対象: Windows x64 / AppWatchdog v1.0.0以降")
    set_run_font(run, size=9.5, color=MUTED)

    doc.add_page_break()


def add_inline_runs(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\[[^\]]+\]\([^)]+\))", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.2, color=DARK_BLUE)
        else:
            link_match = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", part)
            display = link_match.group(1) if link_match else part
            run = paragraph.add_run(display)
            set_run_font(run, size=10)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=8.8, color=RGBColor(40, 40, 40))


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    col_count = len(rows[0])
    widths = TABLE_WIDTHS_DXA.get(col_count, [9360 // col_count] * col_count)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.cell(row_idx, col_idx)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_inline_runs(p, value.replace("**", ""))
            for run in p.runs:
                run.bold = row_idx == 0
                run.font.size = Pt(10 if row_idx == 0 else 9.5)
                if row_idx == 0:
                    run.font.color.rgb = DARK_BLUE
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_numbering_definition(doc: Document, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def parse_markdown(doc: Document, text: str) -> None:
    lines = text.splitlines()
    index = 1  # Skip Markdown title because the DOCX has a cover.
    in_code = False
    code_lines: list[str] = []
    current_list_type: str | None = None
    current_list_num_id: int | None = None

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            current_list_type = None
            current_list_num_id = None
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[-:| ]+\|$", lines[index + 1].strip()):
            current_list_type = None
            current_list_num_id = None
            table_rows = [[cell.strip() for cell in stripped.strip("|").split("|")]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_markdown_table(doc, table_rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            current_list_type = None
            current_list_num_id = None
            level = min(len(heading.group(1)) - 1, 3)
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline_runs(p, heading.group(2))
            index += 1
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if current_list_type != "number":
                current_list_type = "number"
                current_list_num_id = add_numbering_definition(doc, ordered=True)
            p = doc.add_paragraph()
            apply_numbering(p, current_list_num_id)
            add_inline_runs(p, numbered.group(1))
            index += 1
            continue

        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            if current_list_type != "bullet":
                current_list_type = "bullet"
                current_list_num_id = add_numbering_definition(doc, ordered=False)
            p = doc.add_paragraph()
            apply_numbering(p, current_list_num_id)
            add_inline_runs(p, bullet.group(1))
            index += 1
            continue

        current_list_type = None
        current_list_num_id = None
        p = doc.add_paragraph()
        if stripped.endswith(":") or stripped in ("重要:", "注意:"):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
            add_inline_runs(p, stripped)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = DARK_BLUE
        else:
            add_inline_runs(p, stripped.replace("**", ""))
        index += 1


def main() -> int:
    if not SOURCE.exists():
        print(f"Source not found: {SOURCE}", file=sys.stderr)
        return 1

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    add_cover(doc)
    parse_markdown(doc, SOURCE.read_text(encoding="utf-8"))

    doc.core_properties.title = "AppWatchdog Windows 操作マニュアル"
    doc.core_properties.subject = "AppWatchdogのWindows向け利用手順"
    doc.core_properties.author = "AppWatchdog"
    doc.core_properties.comments = "Generated from docs/WINDOWS_USER_MANUAL_JA.md"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
